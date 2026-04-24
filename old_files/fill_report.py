"""
Update Project_Design_Status_Report_Group18.docx with all extracted data.
Saves as Project_Design_Status_Report_Group18_FILLED.docx.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor, Pt
import copy

def set_cell_text(cell, text, bold=False, font_size=None, color=None):
    """Set cell text, preserving first paragraph formatting."""
    para = cell.paragraphs[0]
    # Clear existing runs
    for run in para.runs:
        run.text = ''
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()
    run.text = text
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)

doc = Document('Project_Design_Status_Report_Group18.docx')

# ─────────────────────────────────────────────────────────────
# TABLE 1 — Scope Tracking  (rows 0=header, 1-8 = data)
# Col 2 = Status, Col 3 = Expected Completion, Col 4 = Evidence
# Row 6 = MMIO row (Goal System, MMIO row) — fix address
# ─────────────────────────────────────────────────────────────
t0 = doc.tables[0]

# Fix MMIO address in row 6 (index 6, Goal System MMIO)
# Current text includes 0x80000028 for START — fix to 0x80000024
mmio_row = t0.rows[6]
current_text = mmio_row.cells[0].text
if '0x80000028' in current_text and 'START' in current_text:
    para = mmio_row.cells[0].paragraphs[0]
    for run in para.runs:
        if '0x80000028' in run.text:
            run.text = run.text.replace(
                'MMIO decoder: kernel regs 0x80000000–0x80000020, START at 0x80000028, STATUS at 0x8000002C, with DMEM passthrough.',
                'MMIO decoder: kernel regs 0x80000000–0x80000020, START at 0x80000024, STATUS at 0x80000028, SW_DONE at 0x80000034, NORM at 0x80000030. DMEM passthrough for non-MMIO addresses.'
            )
            break

# Also fix evidence for MMIO row
mmio_row.cells[4].paragraphs[0].runs[0].text = 'Design: 2965104'

# Update Stretch Goal row 7 (sw_gaussian_blur.c) status to Done
t0.rows[7].cells[2].paragraphs[0].runs[0].text = 'Done'

# ─────────────────────────────────────────────────────────────
# TABLE 2 — Design Goals / Performance Metrics
# Col 2 = Current, Col 3 = Status
# ─────────────────────────────────────────────────────────────
t1 = doc.tables[1]
# Row 1 = Latency — TBD (leave as-is, FPGA synthesis not done)
# Row 2 = Throughput → already Met
# Rows 3-6 = Resources → TBD (no synthesis)
# Row 7 = Correctness → already Met (simulation)

# ─────────────────────────────────────────────────────────────
# TABLE 4 — Risk Tracking  (add new rows)
# ─────────────────────────────────────────────────────────────
t4 = doc.tables[4]
# Add row: MMIO address mismatch
new_risks = [
    [
        'MMIO address mismatch — doc says START=0x80000028 but RTL (mmio_decoder.v) uses START=0x80000024.',
        'Ongoing',
        'Medium',
        'Satish / Shaurya',
        'Confirm final address map between RTL and C code (hw_conv_mmio.c) before demo.'
    ],
    [
        'FPGA synthesis and bitstream not yet generated — resource and timing results unknown.',
        'Ongoing',
        'High',
        'Shaurya',
        'Run Vivado implementation before Apr 25; check LUT/FF/BRAM/DSP utilization and timing report.'
    ]
]
for risk_data in new_risks:
    # Copy formatting from last data row
    tr = t4.rows[-1]._tr
    new_tr = copy.deepcopy(tr)
    t4._tbl.append(new_tr)
    new_row = t4.rows[-1]
    for ci, text in enumerate(risk_data):
        cell = new_row.cells[ci]
        para = cell.paragraphs[0]
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)

# ─────────────────────────────────────────────────────────────
# TABLE 5 — Per-Person Contribution
# Update Shaurya integration row status + Sukhjot SW support
# ─────────────────────────────────────────────────────────────
t5 = doc.tables[5]
# Row 3 = Shaurya integration WIP
t5.rows[3].cells[2].paragraphs[0].runs[0].text = 'WIP'
t5.rows[3].cells[3].paragraphs[0].runs[0].text = 'Apr 25, 2026'
# Row 4 = Sukhjot — update to reflect latest SW support work
t5.rows[4].cells[1].paragraphs[0].runs[0].text = (
    'hazard_unit.v (fwd + stall) — Done Apr 5; decoder_ext.v (RV32M decode) — Done Apr 5; '
    'SW branch consolidation and final RTL fixes — Done Apr 19'
)
t5.rows[4].cells[2].paragraphs[0].runs[0].text = 'Done'
t5.rows[4].cells[3].paragraphs[0].runs[0].text = 'Apr 5 / Apr 19, 2026'

# ─────────────────────────────────────────────────────────────
# TABLE 6 — GitHub commit counts (updated from git shortlog)
# ─────────────────────────────────────────────────────────────
t6 = doc.tables[6]

def set_cell_simple(cell, text):
    para = cell.paragraphs[0]
    # Remove all runs first
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

# Sukhjot: 24 (rajpreet.js459) + 7 (Sukhjot-SinghS) = 31
set_cell_simple(t6.rows[2].cells[1], '31')
# Soumik: 6 (roysoumik286)
set_cell_simple(t6.rows[3].cells[1], '6')
# Recalculate total: 12 + 31 + 6 + 27 + 27 = 103
set_cell_simple(t6.rows[6].cells[1], '103')

# ─────────────────────────────────────────────────────────────
# TABLE 7 — Key Commits (add 2 more commits, update existing)
# ─────────────────────────────────────────────────────────────
t7 = doc.tables[7]

# Fix commit 1 title to include hash correctly
t7.rows[1].cells[0].paragraphs[0].runs[0].text = '8b22e70: Integrate coprocessor, UART + SW utilities (Clean)'
t7.rows[1].cells[1].paragraphs[0].runs[0].text = (
    'Full Phase 2 integration: fixed line_buffer BRAM latency via pipeline register insertion, '
    'wired top_fsm 4-cycle drain state, brought in UART controller and C workloads; 131/131 simulation tests passing post-merge.'
)

# Fix commit 2 — update MMIO address description to match RTL
t7.rows[2].cells[0].paragraphs[0].runs[0].text = '2965104: Implement full MMIO decoder for kernel regfile + status reads'
t7.rows[2].cells[1].paragraphs[0].runs[0].text = (
    'Maps kernel weights to 0x80000000–0x80000020 (9 regs), START to 0x80000024, STATUS to 0x80000028, '
    'SW_DONE doorbell at 0x80000034, NORM toggle at 0x80000030; adds DMEM passthrough for non-MMIO addresses.'
)

# Fix commit 3 — GUI
t7.rows[3].cells[0].paragraphs[0].runs[0].text = '87ff069: Add CustomTkinter GUI for FPGA coprocessor interface'
t7.rows[3].cells[1].paragraphs[0].runs[0].text = (
    'Premium dark-mode GUI: sends 128x128 grayscale image over UART, reads 126x126 hardware-convolved result, '
    'displays side-by-side comparison and saves output as timestamped PNG. Uses threaded UART worker to avoid GUI freeze.'
)

# Add 2 more commits
extra_commits = [
    [
        'b163f42: Merge Sukhjots fixes into shaurya branch',
        'Final integration commit: imported all working RTL from Sukhjots branch; resolved 28 file conflicts prioritizing functional fixes; patched Makefiles for Windows (cp -> copy /Y); system runs end-to-end on shaurya branch.'
    ],
    [
        '56357c8: Add RV32M ALU module',
        'Initial rv32m_alu.v commit: implements all 8 RV32M instructions. Multiplies are single-cycle using DSP48E1 inference. Divides use a 32-cycle iterative restoring-division FSM with alu_busy_o stall signal.'
    ],
]
for commit_data in extra_commits:
    tr = t7.rows[-1]._tr
    new_tr = copy.deepcopy(tr)
    t7._tbl.append(new_tr)
    new_row = t7.rows[-1]
    for ci, text in enumerate(commit_data):
        cell = new_row.cells[ci]
        para = cell.paragraphs[0]
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)

# ─────────────────────────────────────────────────────────────
# TABLE 8 — TA Inputs (fill with placeholder rows)
# ─────────────────────────────────────────────────────────────
t8 = doc.tables[8]
ta_note = 'TA feedback to be added after Apr 25-26 final demo review.'
t8.rows[1].cells[0].paragraphs[0].runs[0].text if t8.rows[1].cells[0].paragraphs[0].runs else t8.rows[1].cells[0].paragraphs[0].add_run(ta_note)

# ─────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────
out_path = 'Project_Design_Status_Report_Group18_FILLED.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
